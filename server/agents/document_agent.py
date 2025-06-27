from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentAgent:
    def write_document_file(self,content):
        """Generate a story based on the user's prompt with optional page limit"""
        document = Document()
        
        # Adding a title
        document.add_heading('Comprehensive AI Report', level=0)

        # Process content and apply formatting
        for line in content.strip().split('\n'):
            line = line.strip()
            if line.startswith('#'):
                if line.startswith('##'):
                    document.add_heading(line[3:], level=2)
                else:
                    document.add_heading(line[2:], level=1)
            elif line.startswith('* '):
                p = document.add_paragraph(style='List Bullet')
                # Basic bolding example - more sophisticated parsing needed for complex cases
                parts = line[2:].split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1: # Odd indices are bolded text
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            elif line: # Regular paragraph
                document.add_paragraph(line)
        #document.add_picture('path_to_your_image.jpg', width=Pt(300))
        
        return document


def main():
    chatgpt_content = """
    # Introduction to AI
    Artificial Intelligence (AI) is a rapidly evolving field.
    ### Machine Learning
    Machine learning is a subset of AI that focuses on enabling systems to learn from data. Key concepts include:
    *   **Supervised learning:** Learning from labeled data.
    *   **Unsupervised learning:** Discovering patterns in unlabeled data.
    *   **Reinforcement learning:** Learning through trial and error.
    ### Deep Learning
    Deep learning is a specialized area of machine learning using neural networks.
    """
    try:
        # Get user input
        print("Welcome! I am an Document Publishing Agent ....")
        filename = "Test_Document.docx"
                
        # Initialize the agent
        document_agent = DocumentAgent()
        
        # Get the Content
        print("\nGenerating Content Document...")
        document =  document_agent.write_document_file(chatgpt_content)
        # Save the document
        document.save(filename)
        
        
    except Exception as e:
        print(f"Error generating story: {str(e)}")

if __name__ == "__main__":
    # Run the example
    main()

